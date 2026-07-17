#!/usr/bin/env python3
"""
build_diagnostico_docx.py — Gera o Diagnóstico (Interno OU Cliente) em DOCX.

Recebe um JSON estruturado com o conteúdo do diagnóstico e produz um DOCX
formatado conforme `references/estrutura_diagnostico.md`. Hyperlinks são
inseridos com âncoras "Abrir aula" / "Abrir catálogo" para preservar o padrão
visual do material do Marcelo Coelho.

Uso:
    python build_diagnostico_docx.py --modo [interno|cliente] \\
        --json caminho/diagnostico.json --saida caminho/saida.docx

Estrutura do JSON em EXEMPLO_PAYLOAD no rodapé deste arquivo.

Dependências: python-docx
"""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm


# -------- helpers de estilo ---------------------------------------------------

DARK_TEXT = RGBColor(0x1F, 0x1F, 0x23)
SUBTLE_TEXT = RGBColor(0x6B, 0x72, 0x80)
ACCENT_CYAN = RGBColor(0x06, 0xB6, 0xD4)


def _set_cell_shading(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def _add_hyperlink(paragraph, url: str, text: str, color_hex: str = '06B6D4') -> None:
    """Insere um hyperlink real no parágrafo (não só texto azul)."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True,
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    r_pr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), color_hex)
    r_pr.append(color)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    r_pr.append(u)
    new_run.append(r_pr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _set_default_font(doc: Document, name: str = 'Calibri', size_pt: float = 11) -> None:
    style = doc.styles['Normal']
    style.font.name = name
    style.font.size = Pt(size_pt)


def _heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = DARK_TEXT


def _p(doc: Document, text: str, bold: bool = False, italic: bool = False,
       size: float | None = None, color: RGBColor | None = None,
       align: int | None = None) -> None:
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color


def _bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def _kv_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    """Tabela de 2 colunas (chave/valor) sem cabeçalho."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = 'Light Grid Accent 1'
    for i, (k, v) in enumerate(rows):
        cell_k = table.cell(i, 0)
        cell_v = table.cell(i, 1)
        cell_k.text = ''
        cell_v.text = ''
        run_k = cell_k.paragraphs[0].add_run(k)
        run_k.bold = True
        cell_v.paragraphs[0].add_run(str(v) if v is not None else 'Não informado')


def _matrix_table(doc: Document, header: list[str], rows: list[list[str]],
                  hyperlink_col: int | None = None) -> None:
    """Tabela com cabeçalho. Se hyperlink_col != None, a coluna esperada já contém um dict {'text':..., 'url':...}."""
    table = doc.add_table(rows=len(rows) + 1, cols=len(header))
    table.style = 'Light Grid Accent 1'

    for i, h in enumerate(header):
        cell = table.cell(0, i)
        cell.text = ''
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        _set_cell_shading(cell, 'E5F6FA')

    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ''
            if hyperlink_col is not None and c_idx == hyperlink_col and isinstance(val, dict):
                _add_hyperlink(
                    cell.paragraphs[0],
                    val['url'],
                    val.get('text', 'Abrir'),
                )
            else:
                cell.paragraphs[0].add_run(str(val) if val is not None else '')


# -------- construção do documento --------------------------------------------

def build(modo: str, data: dict, saida: Path) -> None:
    assert modo in ('interno', 'cliente'), 'modo deve ser interno ou cliente'

    doc = Document()
    _set_default_font(doc, 'Calibri', 11)
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    cliente = data['cliente']
    empresa = data.get('empresa', '')

    # Capa
    titulo = (
        'Diagnóstico Interno de Automação Inteligente'
        if modo == 'interno'
        else 'Diagnóstico Inicial de Automação Inteligente'
    )
    _p(doc, titulo, bold=True, size=18, color=DARK_TEXT)
    _p(doc, f'{cliente} - {empresa}', size=13, color=SUBTLE_TEXT)
    if modo == 'interno':
        _p(doc, 'Documento de uso interno para registro, acompanhamento e consulta futura',
           italic=True, size=10, color=SUBTLE_TEXT)
    doc.add_paragraph()

    # 1. Identificação
    _heading(doc, '1. Identificação', level=1)
    ident_rows = [
        ('Cliente', cliente),
        ('Empresa', empresa),
        ('CNPJ', data.get('cnpj', 'Não informado')),
        ('E-mail registrado no chat da call', data.get('email', 'Não informado')),
        ('Setor', data.get('setor', 'Não informado')),
        ('Faturamento anual informado', data.get('faturamento', 'Não informado')),
        ('Folha mensal informada', data.get('folha', 'Não informado')),
        ('Equipe operacional direta', data.get('equipe', 'Não informado')),
        ('Nível de IA', data.get('nivel_ia', 'Não informado')),
    ]
    if modo == 'interno':
        ident_rows.append(('Status do acesso', data.get('status_acesso', 'Não informado')))
    _kv_table(doc, ident_rows)
    doc.add_paragraph()

    # 2. Síntese executiva
    _heading(doc, '2. Síntese executiva', level=1)
    for par in data['sintese']:
        _p(doc, par)
    doc.add_paragraph()

    # 3. Fontes e escopo
    _heading(doc, '3. Fontes e escopo deste diagnóstico', level=1)
    _p(doc, data.get('escopo', ''))
    if data.get('fontes'):
        _matrix_table(doc, ['Fonte', 'Uso no diagnóstico'], data['fontes'])
    doc.add_paragraph()

    # 4. Formulário reconstituído
    _heading(doc, '4. Formulário de diagnóstico reconstituído', level=1)
    if data.get('formulario'):
        _kv_table(doc, data['formulario'])
    doc.add_paragraph()

    # 5. Diagnóstico consultivo consolidado
    _heading(doc, '5. Diagnóstico consultivo consolidado', level=1)
    for par in data['consolidado']:
        _p(doc, par)

    _heading(doc, '5.1 Gargalo central', level=2)
    _bullets(doc, data['gargalo_central'])

    _heading(doc, '5.2 Por que a frente prioritária vem primeiro', level=2)
    for par in data['justificativa_prioridade']:
        _p(doc, par)
    _p(doc, f"Decisão de prioridade: {data['decisao_prioridade']}",
       bold=True, color=ACCENT_CYAN)
    doc.add_paragraph()

    # 6. Frentes de negócio
    _heading(doc, '6. Frentes de negócio', level=1)
    _matrix_table(
        doc,
        ['Frente de negócio', 'Descrição', 'Oportunidade de automação/IA'],
        data['frentes_negocio'],
    )
    doc.add_paragraph()

    # 7. Mapa de atividades por área
    _heading(doc, '7. Mapa de atividades por área e horas/semana', level=1)
    _matrix_table(
        doc,
        ['Atividade', 'Área', 'Horas/semana', 'Potencial de automação'],
        data['atividades'],
    )
    doc.add_paragraph()

    # 8. Resultados esperados
    _heading(doc, '8. Resultados esperados pelo cliente', level=1)
    _matrix_table(
        doc,
        ['Resultados esperados registrados', 'Leitura consultiva'],
        data['resultados_esperados'],
    )
    doc.add_paragraph()

    # 9. Maturidade (só interno)
    if modo == 'interno' and data.get('maturidade'):
        _heading(doc, '9. Maturidade e riscos', level=1)
        _matrix_table(
            doc,
            ['Dimensão', 'Nota (1-5)', 'Justificativa'],
            data['maturidade'],
        )
        doc.add_paragraph()

    # 10. Riscos (só interno)
    if modo == 'interno' and data.get('riscos'):
        _heading(doc, '10. Riscos e como mitigar', level=1)
        _matrix_table(doc, ['Risco', 'Como mitigar'], data['riscos'])
        doc.add_paragraph()

    # 11. Prioridades de implementação
    section_num = 11 if modo == 'interno' else 9
    _heading(doc, f'{section_num}. Prioridades de implementação', level=1)
    _matrix_table(
        doc,
        ['Prioridade', 'Frente', 'Objetivo', 'Entrega esperada'],
        data['prioridades'],
    )
    doc.add_paragraph()

    # 12. Por que a #1 primeiro
    _heading(doc, f'{section_num + 1}. Por que a prioridade nº1 vem primeiro', level=1)
    _matrix_table(doc, ['Motivo', 'Implicação para a mentoria'], data['por_que_primeiro'])
    doc.add_paragraph()

    # 13. Estrutura recomendada do projeto-âncora
    if data.get('estrutura_projeto'):
        _heading(doc, f'{section_num + 2}. Estrutura recomendada do projeto-âncora', level=1)
        for par in data['estrutura_projeto'].get('texto', []):
            _p(doc, par)
        if data['estrutura_projeto'].get('categorias_indicadores'):
            _matrix_table(
                doc,
                ['Categoria', 'Indicadores possíveis', 'Uso no relatório'],
                data['estrutura_projeto']['categorias_indicadores'],
            )
        if data['estrutura_projeto'].get('perguntas_estrategicas'):
            doc.add_paragraph()
            _matrix_table(
                doc,
                ['Pergunta estratégica', 'Por que importa'],
                data['estrutura_projeto']['perguntas_estrategicas'],
            )
        doc.add_paragraph()

    # 15. Trilha de aulas
    _heading(doc, 'Trilha de aulas recomendada', level=1)
    _p(doc,
       'Ordem prioriza o momento atual do aluno. Cada link abre direto a aula no portal.',
       color=SUBTLE_TEXT, size=10)
    aulas = data['trilha_aulas']
    aula_rows = [
        [str(i + 1), a['nome'], a['curso'], a['aplicacao'],
         {'text': 'Abrir aula', 'url': a['link']}]
        for i, a in enumerate(aulas)
    ]
    _matrix_table(
        doc,
        ['Ordem', 'Aula', 'Curso/Módulo', 'Aplicação direta', 'Link'],
        aula_rows,
        hyperlink_col=4,
    )
    doc.add_paragraph()

    # 16. Recursos da plataforma
    _heading(doc, 'Recursos da plataforma recomendados', level=1)
    recursos = data['recursos_plataforma']
    rec_rows = [
        [r['etapa'], r['recurso'], r['categoria'], r['por_que'],
         {'text': r.get('texto_link', 'Abrir catálogo'), 'url': r['link']}]
        for r in recursos
    ]
    _matrix_table(
        doc,
        ['Etapa', 'Recurso', 'Categoria', 'Por que usar', 'Link'],
        rec_rows,
        hyperlink_col=4,
    )
    doc.add_paragraph()

    # 17. Roadmap
    _heading(doc, 'Roadmap 7 / 30 / 60 / 90 dias', level=1)
    _matrix_table(
        doc,
        ['Prazo', 'Foco', 'Ações', 'Entregas'],
        data['roadmap'],
    )
    doc.add_paragraph()

    # 18. Participação nas recorrências
    _heading(doc, 'Participação nas recorrências da mentoria', level=1)
    _p(doc, data.get('recorrencias_intro',
                     'Use as recorrências como ambiente de aceleração prática, não apenas como aulas complementares.'))
    _bullets(doc, data['recorrencias_bullets'])
    doc.add_paragraph()

    # 19-21. Seções só do interno
    if modo == 'interno':
        _heading(doc, 'Recomendações para Eric e acompanhamento interno', level=1)
        _matrix_table(
            doc,
            ['Ponto de atenção', 'Recomendação interna'],
            data['recomendacoes_eric'],
        )
        doc.add_paragraph()

        if data.get('mensagem_envio'):
            _heading(doc, 'Mensagem sugerida para WhatsApp', level=1)
            for par in data['mensagem_envio']:
                _p(doc, par)
            doc.add_paragraph()

        if data.get('roteiro_apresentacao'):
            _heading(doc, 'Roteiro da apresentação ao aluno', level=1)
            for i, step in enumerate(data['roteiro_apresentacao'], start=1):
                doc.add_paragraph(f'{i}. {step}')

    saida.parent.mkdir(parents=True, exist_ok=True)
    doc.save(saida)


def main(argv: list[str]) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument('--modo', required=True, choices=['interno', 'cliente'])
    parser.add_argument('--json', required=True, type=Path)
    parser.add_argument('--saida', required=True, type=Path)
    args = parser.parse_args(argv)

    data = json.loads(args.json.read_text(encoding='utf-8'))
    build(args.modo, data, args.saida)
    print(f'OK: {args.saida}')
    return 0


# Esquema esperado do JSON, com exemplo mínimo
EXEMPLO_PAYLOAD = {
    'cliente': 'Marcelo Coelho',
    'empresa': 'Coelho Music Produções LTDA',
    'cnpj': '52.060.306/0001-37',
    'email': 'coelhomusicprod@gmail.com',
    'setor': 'Educação / EdTech / Educação musical',
    'faturamento': 'Até R$ 1 milhão',
    'folha': 'R$ 50.000',
    'equipe': '1 pessoa operacional',
    'nivel_ia': 'Intermediário',
    'status_acesso': 'Acesso ao portal orientado durante a call.',
    'sintese': [
        'Marcelo Coelho possui um negócio educacional híbrido [...]',
        'O gargalo atual não está na falta de produto [...]',
    ],
    'escopo': 'Este documento consolida informações da call [...]',
    'fontes': [
        ['Call de onboarding', 'Base principal para contexto, dores e prioridades.'],
    ],
    'formulario': [
        ('Faturamento anual', 'Até R$ 1 milhão'),
    ],
    'consolidado': ['[parágrafo 1]', '[parágrafo 2]'],
    'gargalo_central': ['bullet 1', 'bullet 2'],
    'justificativa_prioridade': ['parágrafo'],
    'decisao_prioridade': 'começar pelo dashboard pedagógico.',
    'frentes_negocio': [
        ['Jazz Corporativo', 'Palestra B2B...', 'Captação, proposta, follow-up'],
    ],
    'atividades': [
        ['Responder e-mails', 'Comercial', '3h', 'Alto'],
        ['Total mapeado', '-', '~41h', 'Prioridade de padronização'],
    ],
    'resultados_esperados': [['Aumentar leads', 'Requer funil B2B...']],
    'maturidade': [['Maturidade em IA', '3,5', 'Já usa múltiplas ferramentas...']],
    'riscos': [['Uso disperso de IA', 'Definir arquitetura por função']],
    'prioridades': [
        ['1', 'Dashboard', 'Transformar dashboard em ferramenta de gestão',
         'Estrutura de indicadores e visualizações'],
    ],
    'por_que_primeiro': [['Já está em construção', 'Reduz fricção']],
    'estrutura_projeto': {
        'texto': ['O dashboard deve partir de perguntas estratégicas...'],
        'categorias_indicadores': [
            ['Engajamento', 'Frequência, participação', 'Mostrar adesão'],
        ],
        'perguntas_estrategicas': [
            ['Quais escolas estão mais engajadas?', 'Ajuda a identificar sucesso'],
        ],
    },
    'trilha_aulas': [
        {
            'nome': 'Projetos para construção de relatórios, planilhas, planejamentos',
            'curso': 'Lives e aulas abertas',
            'aplicacao': 'Base inicial para dashboard e relatórios.',
            'link': 'https://areadoaluno.expertintegrado.com.br/app/courses/370695b1-487b-45df-b9fc-07981f6b6e9f/lessons/df7f412a-bf34-4ff8-8186-70df43558a64',
        },
    ],
    'recursos_plataforma': [
        {
            'etapa': 'Agora',
            'recurso': 'Gestor de Métricas',
            'categoria': 'Template Lovable',
            'por_que': 'Base para dashboard, KPIs e relatórios.',
            'link': 'https://areadoaluno.expertintegrado.com.br/app/lovable-catalog',
            'texto_link': 'Abrir catálogo Lovable',
        },
    ],
    'roadmap': [
        ['7 dias', 'Ativação e foco',
         'Registrar diagnóstico, primeira aula, levantar dados.',
         'Diagnóstico registrado, trilha iniciada.'],
    ],
    'recorrencias_bullets': [
        'Levar dúvidas sobre estrutura do dashboard.',
        'Validar indicadores e visualizações.',
    ],
    'recomendacoes_eric': [
        ['Não pedir novo formulário ao cliente', 'Usar respostas reconstituídas.'],
    ],
    'mensagem_envio': ['Marcelo, organizei nosso diagnóstico...'],
    'roteiro_apresentacao': [
        'Abrir reforçando que o diagnóstico foi construído a partir da call.',
        'Mostrar o entendimento do negócio.',
    ],
}


if __name__ == '__main__':
    sys.exit(main(sys.argv[1:]))
